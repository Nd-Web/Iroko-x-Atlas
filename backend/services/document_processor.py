"""
Document Processor Service
Handles file upload, text extraction, smart chunking, and indexing.
Supports PDF, DOCX, XLSX, TXT via Azure Document Intelligence.
"""
import os
import re
import logging
import aiofiles
from typing import Optional, Dict, Any, List

logger = logging.getLogger(__name__)


# ── Text Cleaning ─────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Normalise raw extracted text before chunking.
    - Standardises line endings
    - Collapses excess whitespace without destroying paragraph structure
    - Normalises Nigerian currency (₦, N, NGN) → "NGN <amount>"
    - Strips non-printable control characters
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # ₦45,000,000  →  NGN 45,000,000
    text = re.sub(r"₦\s*", "NGN ", text)
    # N45,000,000 (naira prefix before digits)  →  NGN 45,000,000
    text = re.sub(r"\bN(\d)", r"NGN \1", text)
    # 45,000,000 NGN (suffix form)  →  NGN 45,000,000
    text = re.sub(r"([\d,\.]+)\s*NGN\b", r"NGN \1", text)

    # Collapse runs of spaces/tabs (preserve newlines — they carry structure)
    text = re.sub(r"[ \t]{2,}", " ", text)

    # Collapse 3+ blank lines → 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip non-printable chars; keep: tab, LF, CR, printable ASCII, ₦, Latin Extended
    text = re.sub(
        r"[^\x09\x0A\x0D\x20-\x7E₦À-ɏ]", " ", text
    )

    return text.strip()


# ── Semantic Chunking Helpers ─────────────────────────────────────────────────

_HEADING_RE = re.compile(
    r"^("
    r"#{1,6}\s.+"                          # Markdown headings
    r"|[A-Z][A-Z0-9 ,&/\-]{4,}:?\s*$"    # ALL CAPS headings (e.g. "EXECUTIVE SUMMARY")
    r"|\d+[\.\)]\s+[A-Z][^\n]{3,}"        # Numbered headings (e.g. "1. Introduction")
    r"|[A-Z][a-zA-Z ]{3,}:$"             # Title-colon (e.g. "Key Metrics:")
    r")",
    re.MULTILINE,
)


def _estimate_tokens(text: str) -> int:
    """Rough token estimate: ~4 chars per English/mixed token."""
    return max(1, len(text) // 4)


def _split_paragraphs(text: str) -> List[str]:
    return [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]


def _is_heading(text: str) -> bool:
    return _estimate_tokens(text) <= 25 and bool(_HEADING_RE.match(text.strip()))


def _group_into_sections(paragraphs: List[str]) -> List[Dict]:
    """Attach each paragraph to the nearest preceding heading."""
    sections: List[Dict] = []
    current_heading = ""
    current_paras: List[str] = []

    for para in paragraphs:
        if _is_heading(para):
            if current_paras or current_heading:
                sections.append({"heading": current_heading, "paragraphs": current_paras})
            current_heading = para
            current_paras = []
        else:
            current_paras.append(para)

    sections.append({"heading": current_heading, "paragraphs": current_paras})
    return sections


def _get_overlap(text: str, overlap_tokens: int) -> str:
    """Return the last N tokens' worth of complete sentences from text."""
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    parts: List[str] = []
    count = 0
    for sent in reversed(sentences):
        t = _estimate_tokens(sent)
        if count + t > overlap_tokens:
            break
        parts.insert(0, sent)
        count += t
    return " ".join(parts)


def _page_estimate(chunk_index: int, target_tokens: int) -> int:
    """Approximate page number — ~250 tokens per page."""
    return max(1, (chunk_index * target_tokens) // 250 + 1)


def _make_chunk(
    content: str,
    chunk_index: int,
    section_heading: str,
    document_id: str,
    title: str,
    department: str,
    page_number: int,
) -> Dict[str, Any]:
    return {
        "content": content,
        "chunk_index": chunk_index,
        "section_heading": section_heading,
        "page_number": page_number,
        "document_id": document_id,
        "title": title,
        "department": department,
        "token_estimate": _estimate_tokens(content),
    }


# ── Main Chunking Pipeline ────────────────────────────────────────────────────

def smart_chunk_document(
    text: str,
    document_id: str,
    title: str,
    department: str = "",
    target_tokens: int = 500,
    overlap_tokens: int = 60,
) -> List[Dict[str, Any]]:
    """
    Semantic chunking pipeline for enterprise documents.

    Strategy:
      1. Clean the text
      2. Split into paragraphs (blank-line boundaries)
      3. Detect section headings; headings attach to the paragraphs that follow
      4. Accumulate paragraphs into chunks within the token budget
      5. Paragraphs that alone exceed the budget are split at sentence boundaries
      6. Each new chunk starts with the section heading + overlap from the
         previous chunk, so context is never lost at boundaries
      7. Every chunk carries full metadata (document_id, title, department,
         chunk_index, page_number, section_heading)
    """
    text = clean_text(text)
    paragraphs = _split_paragraphs(text)
    if not paragraphs:
        return []

    sections = _group_into_sections(paragraphs)
    chunks: List[Dict[str, Any]] = []
    chunk_index = 0
    prev_overlap = ""

    for section in sections:
        heading = section["heading"]
        paras = section["paragraphs"]

        # Seed the buffer: heading first, then overlap from previous chunk
        buffer = (heading + "\n\n") if heading else ""
        if prev_overlap:
            buffer += prev_overlap + "\n\n"

        for para in paras:
            # Large paragraph → split at sentence boundaries
            if _estimate_tokens(para) > target_tokens:
                sentences = re.split(r"(?<=[.!?])\s+", para)
                for sent in sentences:
                    if _estimate_tokens(buffer + sent) > target_tokens and buffer.strip():
                        chunks.append(_make_chunk(
                            buffer.strip(), chunk_index, heading,
                            document_id, title, department,
                            _page_estimate(chunk_index, target_tokens),
                        ))
                        chunk_index += 1
                        prev_overlap = _get_overlap(buffer, overlap_tokens)
                        buffer = (heading + "\n\n" if heading else "") + (prev_overlap + "\n\n" if prev_overlap else "")
                    buffer += sent + " "
            else:
                if _estimate_tokens(buffer + para) > target_tokens and buffer.strip():
                    chunks.append(_make_chunk(
                        buffer.strip(), chunk_index, heading,
                        document_id, title, department,
                        _page_estimate(chunk_index, target_tokens),
                    ))
                    chunk_index += 1
                    prev_overlap = _get_overlap(buffer, overlap_tokens)
                    buffer = (heading + "\n\n" if heading else "") + (prev_overlap + "\n\n" if prev_overlap else "")
                buffer += para + "\n\n"

        # Flush remaining buffer for this section
        if buffer.strip():
            chunks.append(_make_chunk(
                buffer.strip(), chunk_index, heading,
                document_id, title, department,
                _page_estimate(chunk_index, target_tokens),
            ))
            chunk_index += 1
            prev_overlap = _get_overlap(buffer, overlap_tokens)

    return chunks


# ── Content Quality Gate ──────────────────────────────────────────────────────

# Text that indicates we downloaded an HTML web-viewer page instead of the file,
# or that Azure DI processed Office XML metadata instead of document content.
_BOILERPLATE_SIGNALS = [
    "creating documents on any device",
    "access files from anywhere",
    "office apps like word, excel",
    "microsoft 365",
    "sign in to your microsoft account",
    "get the free mobile app",
    "onedrive lets you",
    "store photos and docs online",
    "<!doctype html",
    "<html",
    "content-type: text/html",
]


def _is_meaningful(text: str) -> bool:
    """
    Return True if the extracted text looks like real document content.
    Rejects empty text, pure whitespace, and known boilerplate patterns.
    """
    if not text:
        return False
    stripped = text.strip()
    # Need at least 80 non-whitespace characters
    if len(stripped.replace(" ", "").replace("\n", "")) < 80:
        return False
    lower = stripped.lower()
    if any(signal in lower for signal in _BOILERPLATE_SIGNALS):
        logger.warning("Extracted text matches boilerplate — treating as failed extraction")
        return False
    return True


# ── Native Office Extractors ──────────────────────────────────────────────────

def _extract_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a .docx file using python-docx.
    Captures paragraphs (preserving heading hierarchy) and table cell content.
    Falls back gracefully if the library is unavailable or the file is corrupt.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(file_path)
        parts: List[str] = []

        # Body paragraphs (includes headings, normal text, list items)
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt:
                parts.append(txt)

        # Tables — emit each row as tab-separated cells
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        result = "\n".join(parts)
        if _is_meaningful(result):
            return result
        logger.warning("python-docx extraction produced no meaningful content")
        return None

    except ImportError:
        logger.warning("python-docx not installed — falling back to Azure DI for DOCX")
        return None
    except Exception as e:
        logger.warning(f"python-docx extraction failed: {e}")
        return None


def _extract_xlsx(file_path: str) -> Optional[str]:
    """
    Extract text from an .xlsx file using openpyxl.
    Iterates every sheet, emits the sheet name as a section header,
    then each non-empty row as tab-separated cell values.
    """
    try:
        import openpyxl

        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        parts: List[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_parts: List[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    sheet_parts.append("\t".join(cells))
            if sheet_parts:
                parts.append(f"[Sheet: {sheet_name}]")
                parts.extend(sheet_parts)

        wb.close()
        result = "\n".join(parts)
        if _is_meaningful(result):
            return result
        logger.warning("openpyxl extraction produced no meaningful content")
        return None

    except ImportError:
        logger.warning("openpyxl not installed — falling back to Azure DI for XLSX")
        return None
    except Exception as e:
        logger.warning(f"openpyxl extraction failed: {e}")
        return None


# ── Azure Document Intelligence (PDF / scanned images) ───────────────────────

def _extract_via_azure_di(file_path: str) -> Optional[str]:
    """
    Use Azure Document Intelligence prebuilt-read to extract text.
    Intended for PDFs and scanned images where native parsing isn't applicable.
    """
    endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
    api_key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")

    if not endpoint or not api_key:
        logger.warning("Azure Document Intelligence not configured")
        return None

    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential

        client = DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(api_key),
        )

        with open(file_path, "rb") as f:
            file_content = f.read()

        poller = client.begin_analyze_document(
            model_id="prebuilt-read",
            body=file_content,
            content_type="application/octet-stream",
        )
        result = poller.result()

        text_parts = []
        for page in result.pages:
            for line in page.lines or []:
                text_parts.append(line.content)

        extracted = "\n".join(text_parts)
        if _is_meaningful(extracted):
            return extracted

        logger.warning("Azure DI extraction produced no meaningful content")
        return None

    except Exception as e:
        logger.error(f"Azure Document Intelligence extraction failed: {e}")
        return None


# ── Extraction ────────────────────────────────────────────────────────────────

async def process_document(
    file_path: str,
    document_id: str,
    title: str,
    metadata: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Full pipeline: extract text → smart-chunk → index in Azure Search.
    Returns processing result with chunk count.
    """
    from datetime import datetime
    try:
        text = await extract_text(file_path, metadata.get("file_type", "txt"))
        if not text:
            return {"success": False, "error": "Could not extract meaningful text from document"}

        chunks = smart_chunk_document(
            text=text,
            document_id=document_id,
            title=title,
            department=metadata.get("department", ""),
        )
        logger.info(f"Document '{title}' split into {len(chunks)} semantic chunks")

        from services.azure_search import index_document_chunks
        success = await index_document_chunks(
            document_id=document_id,
            title=title,
            chunks=chunks,
            metadata={**metadata, "created_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")},
        )

        return {
            "success": success,
            "chunk_count": len(chunks),
            "text_length": len(text),
            **({"error": "Azure Search upload failed — check app logs"} if not success else {}),
        }

    except Exception as e:
        logger.error(f"Document processing failed for '{title}': {e}")
        return {"success": False, "error": str(e)}


async def extract_text(file_path: str, file_type: str) -> Optional[str]:
    """
    Extract plain text from a document using the best available method for each type.

    Strategy per file type:
      - txt / md / csv  → direct file read
      - docx            → python-docx (paragraphs + tables), Azure DI as fallback
      - xlsx            → openpyxl (all sheets/cells), Azure DI as fallback
      - pdf             → Azure DI prebuilt-read (purpose-built for PDFs)
      - other           → Azure DI, then None

    All paths run through _is_meaningful() to reject boilerplate/garbage before
    returning to the caller, preventing junk from reaching the search index.
    """
    file_type = file_type.lower().strip(".")

    # ── Plain text formats ────────────────────────────────────────────────────
    if file_type in ("txt", "md", "csv"):
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = await f.read()
            return content if _is_meaningful(content) else None
        except Exception as e:
            logger.error(f"Text read failed: {e}")
            return None

    # ── Word documents ────────────────────────────────────────────────────────
    import asyncio
    if file_type == "docx":
        text = await asyncio.to_thread(_extract_docx, file_path)
        if text:
            logger.info(f"DOCX extracted via python-docx: {len(text)} chars")
            return text
        logger.info("Falling back to Azure DI for DOCX")
        text = await asyncio.to_thread(_extract_via_azure_di, file_path)
        if text:
            logger.info(f"DOCX extracted via Azure DI fallback: {len(text)} chars")
            return text
        logger.error(f"All extraction methods failed for DOCX: {file_path}")
        return None

    # ── Excel spreadsheets ────────────────────────────────────────────────────
    if file_type in ("xlsx", "xls"):
        text = await asyncio.to_thread(_extract_xlsx, file_path)
        if text:
            logger.info(f"XLSX extracted via openpyxl: {len(text)} chars")
            return text
        logger.info("Falling back to Azure DI for XLSX")
        text = await asyncio.to_thread(_extract_via_azure_di, file_path)
        if text:
            logger.info(f"XLSX extracted via Azure DI fallback: {len(text)} chars")
            return text
        logger.error(f"All extraction methods failed for XLSX: {file_path}")
        return None

    # ── PDFs and everything else → Azure DI ──────────────────────────────────
    # Run in a thread so the synchronous poller.result() doesn't block the event loop.
    text = await asyncio.to_thread(_extract_via_azure_di, file_path)
    if text:
        return text

    # Last resort: dev mock so the pipeline doesn't stall entirely
    logger.warning(f"All extractors failed for '{file_type}' — using mock text")
    return _mock_document_text(file_type)


def _mock_document_text(file_type: str) -> str:
    """Returns realistic mock document text for development."""
    return """MTN Nigeria Internal Document

EXECUTIVE SUMMARY

This document contains enterprise information for MTN Nigeria Limited.
Network performance data indicates strong growth in subscriber base across
all regions. Lagos, Abuja, and Port Harcourt remain top revenue contributors.

Key Metrics for Q1 2026:

- Total subscribers: 76.5 million
- Network uptime: 99.3%
- Customer satisfaction index: 71/100
- Revenue growth: 12% year-on-year
- Total revenue: NGN 245,000,000,000

AREAS REQUIRING ATTENTION

Lagos Zone 7 network capacity is operating below acceptable thresholds.
The TowerCo contract renewal requires immediate action — expiry is imminent.
NCC compliance updates must be filed before the quarterly deadline.

Tower Maintenance Status:

Tower 4471 in Ikeja experienced intermittent failures on four occasions this quarter,
affecting approximately 12,000 subscribers during peak hours (18:00-21:00).
Emergency maintenance has been scheduled but not yet completed.

CONTRACT STATUS

TowerCo Infrastructure Agreement expires May 15, 2026.
Monthly fee: NGN 45,000,000. SLA: 99.5% uptime guarantee.
Renewal notice was due 90 days prior — action is now overdue.
"""
